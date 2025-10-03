"""
Title: Trip Energy Demand Estimation – Random Forest (RF) Model
Author: Jagpreet Singh
Date: 2025-10-03

Description:
    This script trains and evaluates a Random Forest regression model 
    to estimate trip energy demand (Trip_Energy_kw) of battery electric buses 
    in urban public transport systems.

Key Steps:
    1. Load and preprocess the dataset (Dataset.xlsx).
    2. Standardize features and split into train/test sets.
    3. Build and train the Random Forest model with tuned hyperparameters:
       n_estimators=100, max_depth=10, min_samples_split=5, min_samples_leaf=3,
       max_features='sqrt', random_state=42, oob_score=True
    4. Evaluate model performance using R², MAE, MSE, RMSE.
    5. Compute feature importance using SHAP values.
    6. Generate residual distribution and parity plots.
    7. Perform 5-Fold cross-validation and summarize metrics.
    8. Save a comprehensive Word report including metrics, SHAP values, 
       residuals, parity plots, and hyperparameters.

Input:
    - Dataset.xlsx: Contains features
      ['Passengers', 'Speed', 'Route_Length_km', 'Number_of_Bus_Stops',
       'Trip_Duration_Minutes', 'Trip_Temprature'] and target 'Trip_Energy_kw'.

Output:
    - Trained Random Forest model
    - Evaluation metrics (train-test and 5-Fold CV)
    - Feature importance (SHAP)
    - Residual and parity plots (PNG)
    - Word report summarizing model performance and details
"""

import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
import time
import datetime
import shap
from docx import Document
from docx.shared import Inches
from sklearn.model_selection import train_test_split, KFold
from sklearn.preprocessing import StandardScaler
from sklearn.ensemble import RandomForestRegressor
from sklearn.metrics import mean_squared_error, r2_score, mean_absolute_error

# Load the dataset
ds = pd.read_excel("Dataset.xlsx")

# Define input features (X) and target variable (y)
feature_names = ['Passengers', 'Speed', 'Route_Length_km', 'Number_of_Bus_Stops', 'Trip_Duration_Minutes', 'Trip_Temprature']
X = ds[feature_names]
y = ds['Trip_Energy_kw']

# Split the dataset into training and testing sets (80% training, 20% testing)
X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2, random_state=42)

# Scale the input features (Standardization)
scaler = StandardScaler()
X_train_scaled = scaler.fit_transform(X_train)
X_test_scaled = scaler.transform(X_test)

# --- K-Fold Validation (5-fold) ---
kf = KFold(n_splits=5, shuffle=True, random_state=42)
r2_list, mae_list, mse_list, rmse_list = [], [], [], []

for tr_idx, val_idx in kf.split(X_train_scaled):
    # Fit scaler *within each fold* to avoid leakage
    fold_scaler = StandardScaler()
    X_tr = fold_scaler.fit_transform(X_train.iloc[tr_idx])
    X_val = fold_scaler.transform(X_train.iloc[val_idx])
    y_tr = y_train.iloc[tr_idx]
    y_val = y_train.iloc[val_idx]

    rf_cv = RandomForestRegressor(
    n_estimators=100,
    max_depth=10,
    min_samples_split=5,     # avoid splits on very few samples
    min_samples_leaf=3,      # prevent leaves with only 1 sample
    max_features="sqrt",     # increases diversity among trees
    random_state=42,
    oob_score=True           # enables out-of-bag validation
)
    rf_cv.fit(X_tr, y_tr)
    y_val_pred = rf_cv.predict(X_val)

    r2_list.append(r2_score(y_val, y_val_pred))
    mae_list.append(mean_absolute_error(y_val, y_val_pred))
    mse_fold = mean_squared_error(y_val, y_val_pred)
    mse_list.append(mse_fold)
    rmse_list.append(np.sqrt(mse_fold))

# Create and Train Random Forest Regressor Model on full training data (as in original code)
rf_model = RandomForestRegressor(n_estimators=100, max_depth=10, random_state=42)
start_train = time.time()
rf_model.fit(X_train_scaled, y_train)
end_train = time.time()
training_time = end_train - start_train

# Predict on Test Data
start_test = time.time()
y_pred = rf_model.predict(X_test_scaled)
end_test = time.time()
testing_time = end_test - start_test

# Evaluate the Model
mae = mean_absolute_error(y_test, y_pred)
mse = mean_squared_error(y_test, y_pred)
rmse = np.sqrt(mse)
r2 = r2_score(y_test, y_pred)

# Feature Importance using SHAP
explainer = shap.Explainer(rf_model, X_train_scaled)
shap_values = explainer(X_test_scaled)

# Compute mean absolute SHAP values
feature_importance = np.abs(shap_values.values).mean(axis=0)
shap_importance_df = pd.DataFrame({'Feature': feature_names, 'SHAP Importance': feature_importance})
shap_importance_df = shap_importance_df.sort_values(by='SHAP Importance', ascending=False)

# Residual Distribution Plot
residuals = y_test - y_pred
#------------------------------------------------------------------------------------------
# --- Residual Distribution Plot (KDE) with Technical Details ---
plt.figure(figsize=(10, 6), dpi=300)
sns.kdeplot(residuals, fill=True, color="#1a80bb")
sns.despine(top=True, right=True)
plt.grid(True, which='both', linestyle='--', linewidth=0.5, alpha=0.7)
plt.xlabel("Residuals", fontsize=14, fontname='Times New Roman')
plt.ylabel("Density", fontsize=14, fontname='Times New Roman')

# --- Compute Residual Diagnostics ---
mean_resid = residuals.mean()  # Bias check
var_ratio = residuals.var() / y_test.var()  # Homoscedasticity proxy
skewness = residuals.skew()  # Normality check

# --- Add diagnostics inside the plot ---
# resid_text = (
#     f"Mean Residual = {mean_resid:.3f}\n"
#     f"Homoscedasticity = {var_ratio:.3f}\n"
#     f"Skewness = {skewness:.3f}"
# )
# plt.text(
#     0.05, 0.95, resid_text,
#     transform=plt.gca().transAxes,
#     fontname='Times New Roman',
#     fontsize=11,
#     verticalalignment='top',
#     bbox=dict(boxstyle="round,pad=0.3", edgecolor="gray", facecolor="white", alpha=0.8)
# )

# Save residual plot
residual_plot_filename = "Residual_RF.png"
plt.savefig(residual_plot_filename, dpi=300, bbox_inches='tight')
plt.close()

# Save Outputs to a Word Document
doc = Document()
doc.add_heading("Random Forest Model", level=1)
doc.add_paragraph(f"Date: {datetime.datetime.now().strftime('%d-%m-%Y %H:%M')}\n\n")

# Model Performance Metrics
doc.add_heading("1️⃣ Model Performance Metrics", level=2)
table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Metric'
hdr_cells[1].text = 'Value'

metrics = {
    "R-Squared (R2)": r2,
    "Mean Absolute Error (MAE)": mae,
    "Mean Squared Error (MSE)": mse,
    "Root Mean Squared Error (RMSE)": rmse,
    "Training Time (seconds)": training_time,
    "Testing Time (seconds)": testing_time
}
for metric, value in metrics.items():
    row_cells = table.add_row().cells
    row_cells[0].text = metric
    row_cells[1].text = str(round(value, 4))

# 2️⃣ SHAP Feature Importance
doc.add_heading("2️⃣ SHAP Feature Importance", level=2)
table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Feature'
hdr_cells[1].text = 'SHAP Importance'
for _, row in shap_importance_df.iterrows():
    row_cells = table.add_row().cells
    row_cells[0].text = row['Feature']
    row_cells[1].text = str(round(row['SHAP Importance'], 4))

# 3️⃣ K-Fold Validation Metrics (5-Fold)
doc.add_heading("3️⃣ K-Fold Validation Metrics (5-Fold)", level=2)
table = doc.add_table(rows=1, cols=5); table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Model'
hdr_cells[1].text = 'R² (mean ± std)'
hdr_cells[2].text = 'MAE (mean ± std)'
hdr_cells[3].text = 'MSE (mean ± std)'
hdr_cells[4].text = 'RMSE (mean ± std)'

row_cells = table.add_row().cells
row_cells[0].text = 'Random Forest'
row_cells[1].text = f"{np.mean(r2_list):.4f} ± {np.std(r2_list):.4f}"
row_cells[2].text = f"{np.mean(mae_list):.4f} ± {np.std(mae_list):.4f}"
row_cells[3].text = f"{np.mean(mse_list):.4f} ± {np.std(mse_list):.4f}"
row_cells[4].text = f"{np.mean(rmse_list):.4f} ± {np.std(rmse_list):.4f}"

# Append Residual KDE Plot (kept exactly as original)
doc.add_heading("4️⃣ Residual Distribution Plot (KDE) for Random Forest", level=2)
doc.add_picture(residual_plot_filename, width=Inches(5.5))

# Residual Diagnostics Table (kept)
doc.add_heading("Residual Diagnostics", level=3)
table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Diagnostic'
hdr_cells[1].text = 'Value'
diagnostics = {"Mean Residual": mean_resid, "Homoscedasticity": var_ratio, "Skewness": skewness}
for diag, value in diagnostics.items():
    row_cells = table.add_row().cells
    row_cells[0].text = diag
    row_cells[1].text = str(round(value, 4))

# --- Parity Plot (Predicted vs Observed) --------------------------------
plt.figure(figsize=(10, 6), dpi=300)
sns.scatterplot(x=y_test, y=y_pred, alpha=0.5, color="#ffb255", edgecolor=None)
plt.plot([y_test.min(), y_test.max()], [y_test.min(), y_test.max()],
         linestyle="--", color="#000000", lw=2)
sns.despine(top=True, right=True)
plt.grid(True, which="both", linestyle="--", linewidth=0.5, alpha=0.6)
plt.xlabel("Observed Trip Energy (kWh)", fontsize=14, fontname='Times New Roman')
plt.ylabel("Predicted Trip Energy (kWh)", fontsize=14, fontname='Times New Roman')

# Add model metrics inside the plot
# metrics_text = (
#     f"R² = {r2:.3f}\n"
#     f"MSE = {mse:.2f}\n"
#     f"RMSE = {rmse:.2f}\n"
#     f"MAE = {mae:.2f}"
# )
# plt.text(
#     0.05, 0.95, metrics_text,
#     transform=plt.gca().transAxes,
#     fontname='Times New Roman',
#     fontsize=11,
#     verticalalignment='top',
#     bbox=dict(boxstyle="round,pad=0.3", edgecolor="gray", facecolor="white", alpha=0.8)
# )

# Save parity plot
parity_plot_filename = "Parity_RF.png"
plt.savefig(parity_plot_filename, dpi=300, bbox_inches='tight')
plt.close()

# Append Parity Plot to Word Document (kept)
doc.add_heading("5️⃣ Parity Plot (Predicted vs Observed)", level=2)
doc.add_picture(parity_plot_filename, width=Inches(5.5))

# 6️⃣ Model Hyperparameters
doc.add_heading("6️⃣ Model Hyperparameters", level=2)
table = doc.add_table(rows=1, cols=2); table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Model'
hdr_cells[1].text = 'Key Hyperparameters'
row_cells = table.add_row().cells
row_cells[0].text = 'Random Forest'
# Use get_params() for full transparency; or provide a concise string
row_cells[1].text = str(rf_model.get_params())

# --- END Parity Plot ------------------------------------------------
# Save the document
doc.save("2025_Report_RF_KFold_Full.docx")
print("\n✅ Report saved successfully as '2025_Report_RF_KFold_Full.docx'")
