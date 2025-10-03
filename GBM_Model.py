"""
Title: Trip Energy Demand Estimation – Gradient Boosting Machine (GBM) Model
Author: Jagpreet Singh
Date: 2025-10-03

Description:
    This script trains and evaluates a Gradient Boosting Machine (GBM) model 
    to estimate trip energy demand (Trip_Energy_kw) of battery electric buses 
    in urban public transport systems.

Key Steps:
    1. Load and preprocess the dataset (Dataset.xlsx).
    2. Standardize features and split into train/test sets.
    3. Build and train the GBM model with tuned hyperparameters:
       max_depth=2, n_estimators=100, learning_rate=0.05, subsample=0.7, max_features=0.7
    4. Evaluate model performance using R², MAE, MSE, RMSE.
    5. Compute feature importance using SHAP values.
    6. Extract tree statistics (total nodes, decision nodes, leaf nodes).
    7. Generate residual distribution and parity plots.
    8. Perform 5-Fold cross-validation and summarize metrics.
    9. Save a comprehensive Word report including metrics, SHAP values, 
       residuals, parity plots, tree statistics, and hyperparameters.

Input:
    - Dataset.xlsx: Contains features
      ['Passengers', 'Speed', 'Route_Length_km', 'Number_of_Bus_Stops',
       'Trip_Duration_Minutes', 'Trip_Temprature'] and target 'Trip_Energy_kw'.

Output:
    - Trained GBM model
    - Evaluation metrics (train-test and 5-Fold CV)
    - Feature importance (SHAP)
    - Tree statistics
    - Residual and parity plots (PNG)
    - Word report summarizing model performance and details
"""

import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
import shap
import time
import datetime
from docx import Document
from docx.shared import Inches
from sklearn.model_selection import train_test_split, KFold
from sklearn.preprocessing import StandardScaler
from sklearn.ensemble import GradientBoostingRegressor
from sklearn.metrics import mean_squared_error, r2_score, mean_absolute_error
import numpy as np

# Load the dataset
ds = pd.read_excel("Dataset.xlsx")

# Define input features and target
feature_names = ['Passengers', 'Speed', 'Route_Length_km', 'Number_of_Bus_Stops', 'Trip_Duration_Minutes', 'Trip_Temprature']
X = ds[feature_names]
y = ds['Trip_Energy_kw']

# Split dataset for normal training/testing
X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2, random_state=42)

# Scale features
scaler = StandardScaler()
X_train_scaled = scaler.fit_transform(X_train)
X_test_scaled = scaler.transform(X_test)

# ------------------ Train GBM Model ------------------
gbm_model = GradientBoostingRegressor(
    random_state=42,
    max_depth=2,
    n_estimators=100,
    learning_rate=0.05,
    min_samples_split=20,
    min_samples_leaf=10,
    subsample=0.7,
    max_features=0.7
)

start_train = time.time()
gbm_model.fit(X_train_scaled, y_train)
end_train = time.time()
training_time = end_train - start_train

# Predict test data
start_test = time.time()
y_pred = gbm_model.predict(X_test_scaled)
end_test = time.time()
testing_time = end_test - start_test

# ------------------ Evaluate Model ------------------
mse = mean_squared_error(y_test, y_pred)
r2 = r2_score(y_test, y_pred)
mae = mean_absolute_error(y_test, y_pred)
rmse = mean_squared_error(y_test, y_pred, squared=False)

# Tree statistics
total_nodes = sum(estimator.tree_.node_count for estimator in gbm_model.estimators_.flatten())
decision_nodes = sum((estimator.tree_.node_count - estimator.tree_.n_leaves) for estimator in gbm_model.estimators_.flatten())
leaf_nodes = sum(estimator.tree_.n_leaves for estimator in gbm_model.estimators_.flatten())

# SHAP Feature Importance
explainer = shap.Explainer(gbm_model, X_train_scaled)
shap_values = explainer(X_test_scaled)
mean_abs_shap_values = abs(shap_values.values).mean(axis=0)
importance_df = pd.DataFrame({'Feature': feature_names, 'Mean Absolute SHAP Value': mean_abs_shap_values}).sort_values(by='Mean Absolute SHAP Value', ascending=False)

# ------------------ Residual Distribution Plot ------------------
residuals = y_test - y_pred
plt.figure(figsize=(10, 6), dpi=300)
sns.kdeplot(residuals, fill=True, color="#1a80bb")
sns.despine(top=True, right=True)
plt.grid(True, which='both', linestyle='--', linewidth=0.5, alpha=0.7)
plt.xlabel("Residuals", fontsize=14, fontname='Times New Roman')
plt.ylabel("Density", fontsize=14, fontname='Times New Roman')
mean_resid = residuals.mean()
var_ratio = residuals.var() / y_test.var()
skewness = residuals.skew()
resid_text = f"Mean Residual = {mean_resid:.3f}\nHomoscedasticity = {var_ratio:.3f}\nSkewness = {skewness:.3f}"
#plt.text(0.05, 0.95, resid_text, transform=plt.gca().transAxes, fontname='Times New Roman', fontsize=11, verticalalignment='top', bbox=dict(boxstyle="round,pad=0.3", edgecolor="gray", facecolor="white", alpha=0.8))
residual_plot_filename = "Residual_GBM.png"
plt.savefig(residual_plot_filename, dpi=300, bbox_inches='tight')
plt.close()

# ------------------ Parity Plot ------------------
plt.figure(figsize=(10, 6), dpi=300)
sns.scatterplot(x=y_test, y=y_pred, alpha=0.5, color="#ffb255", edgecolor=None)
plt.plot([y_test.min(), y_test.max()], [y_test.min(), y_test.max()], linestyle="--", color="#000000", lw=2)
sns.despine(top=True, right=True)
plt.grid(True, which="both", linestyle="--", linewidth=0.5, alpha=0.6)
plt.xlabel("Observed Trip Energy (kWh)", fontsize=14, fontname='Times New Roman')
plt.ylabel("Predicted Trip Energy (kWh)", fontsize=14, fontname='Times New Roman')
metrics_text = f"R² = {r2:.3f}\nMSE = {mse:.2f}\nRMSE = {rmse:.2f}\nMAE = {mae:.2f}"
#plt.text(0.05, 0.95, metrics_text, transform=plt.gca().transAxes, fontname='Times New Roman', fontsize=11, verticalalignment='top', bbox=dict(boxstyle="round,pad=0.3", edgecolor="gray", facecolor="white", alpha=0.8))
parity_plot_filename = "Parity_GBM.png"
plt.savefig(parity_plot_filename, dpi=300, bbox_inches='tight')
plt.close()

# ------------------ K-Fold Validation ------------------
kf = KFold(n_splits=5, shuffle=True, random_state=42)
r2_list, mae_list, mse_list, rmse_list = [], [], [], []
X_scaled = scaler.fit_transform(X)
for train_index, test_index in kf.split(X_scaled):
    X_tr, X_te = X_scaled[train_index], X_scaled[test_index]
    y_tr, y_te = y.iloc[train_index], y.iloc[test_index]
    model_kf = GradientBoostingRegressor(
        random_state=42,
        max_depth=2,
        n_estimators=100,
        learning_rate=0.05,
        min_samples_split=20,
        min_samples_leaf=10,
        subsample=0.7,
        max_features=0.7
    )
    model_kf.fit(X_tr, y_tr)
    y_pred_kf = model_kf.predict(X_te)
    r2_list.append(r2_score(y_te, y_pred_kf))
    mae_list.append(mean_absolute_error(y_te, y_pred_kf))
    mse_list.append(mean_squared_error(y_te, y_pred_kf))
    rmse_list.append(mean_squared_error(y_te, y_pred_kf, squared=False))

# ------------------ Save Everything to Word ------------------
doc = Document()
doc.add_heading("Gradient Boosting Machine Model", level=1)
doc.add_paragraph(f"Date: {datetime.datetime.now().strftime('%d-%m-%Y %H:%M')}\n\n")

# --- Model Metrics Table ---
doc.add_heading("1️⃣ Model Performance Metrics", level=2)
table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Metric'; hdr_cells[1].text = 'Value'
metrics = {
    "R-Squared (R2)": r2,
    "Mean Absolute Error (MAE)": mae,
    "Mean Squared Error (MSE)": mse,
    "Root Mean Squared Error (RMSE)": rmse,
    "Training Time (seconds)": training_time,
    "Testing Time (seconds)": testing_time
}
for metric, value in metrics.items():
    row_cells = table.add_row().cells
    row_cells[0].text = metric
    row_cells[1].text = str(round(value, 4))

# --- Tree Stats Table ---
doc.add_heading("2️⃣ Tree Statistics", level=2)
table = doc.add_table(rows=1, cols=2); table.style = 'Table Grid'
hdr_cells = table.rows[0].cells; hdr_cells[0].text = 'Statistic'; hdr_cells[1].text = 'Value'
tree_stats = {"Total Number of Nodes": total_nodes, "Number of Decision Nodes": decision_nodes, "Number of Leaf Nodes": leaf_nodes}
for stat, value in tree_stats.items():
    row_cells = table.add_row().cells
    row_cells[0].text = stat; row_cells[1].text = str(value)

# --- Feature Importance ---
doc.add_heading("3️⃣ Feature Importance (Mean Absolute SHAP Values)", level=2)
table = doc.add_table(rows=1, cols=2); table.style = 'Table Grid'
hdr_cells = table.rows[0].cells; hdr_cells[0].text='Feature'; hdr_cells[1].text='Mean Absolute SHAP Value'
for _, row in importance_df.iterrows():
    row_cells = table.add_row().cells
    row_cells[0].text=row['Feature']; row_cells[1].text=str(round(row['Mean Absolute SHAP Value'],4))

# --- Residual Plot ---
doc.add_heading("4️⃣ Residual Distribution Plot (KDE)", level=2)
doc.add_picture(residual_plot_filename, width=Inches(5.5))

# --- Residual Diagnostics ---
doc.add_heading("Residual Diagnostics", level=3)
table = doc.add_table(rows=1, cols=2); table.style='Table Grid'
hdr_cells=table.rows[0].cells; hdr_cells[0].text='Diagnostic'; hdr_cells[1].text='Value'
diagnostics = {"Mean Residual": mean_resid, "Homoscedasticity": var_ratio, "Skewness": skewness}
for diag, value in diagnostics.items():
    row_cells = table.add_row().cells; row_cells[0].text=diag; row_cells[1].text=str(round(value,4))

# --- Parity Plot ---
doc.add_heading("5️⃣ Parity Plot (Predicted vs Observed)", level=2)
doc.add_picture(parity_plot_filename, width=Inches(5.5))

# --- K-Fold Validation Table ---
doc.add_heading("6️⃣ K-Fold Validation Metrics (5-Fold)", level=2)
table = doc.add_table(rows=1, cols=5); table.style='Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text='Model'; hdr_cells[1].text='R² (mean ± std)'; hdr_cells[2].text='MAE (mean ± std)'; hdr_cells[3].text='MSE (mean ± std)'; hdr_cells[4].text='RMSE (mean ± std)'
row_cells = table.add_row().cells
row_cells[0].text='Gradient Boosting Machine'
row_cells[1].text=f"{np.mean(r2_list):.4f} ± {np.std(r2_list):.4f}"
row_cells[2].text=f"{np.mean(mae_list):.4f} ± {np.std(mae_list):.4f}"
row_cells[3].text=f"{np.mean(mse_list):.4f} ± {np.std(mse_list):.4f}"
row_cells[4].text=f"{np.mean(rmse_list):.4f} ± {np.std(rmse_list):.4f}"

# --- Hyperparameters Table ---
doc.add_heading("7️⃣ Model Hyperparameters", level=2)
table = doc.add_table(rows=1, cols=2); table.style='Table Grid'
hdr_cells=table.rows[0].cells; hdr_cells[0].text='Model'; hdr_cells[1].text='Key Hyperparameters'
row_cells = table.add_row().cells
row_cells[0].text='Gradient Boosting Machine'
row_cells[1].text="random_state=42, max_depth=2, n_estimators=100, learning_rate=0.05, min_samples_split=20, min_samples_leaf=10, subsample=0.7, max_features=0.7"

# --- Save Document ---
doc.save("2025_Report_GBM_KFold_Full.docx")
print("\n✅ Report saved successfully as '2025_Report_GBM_KFold_Full.docx'")
