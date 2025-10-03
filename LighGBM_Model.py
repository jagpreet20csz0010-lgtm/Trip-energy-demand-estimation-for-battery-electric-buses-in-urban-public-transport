"""
Title: Trip Energy Demand Estimation – LightGBM Model
Author: Jagpreet Singh
Date: 2025-10-03

Description:
    This script trains and evaluates a LightGBM regression model 
    to estimate trip energy demand (Trip_Energy_kw) of battery electric buses 
    in urban public transport systems.

Key Steps:
    1. Load and preprocess the dataset (Dataset.xlsx).
    2. Standardize features and split into train/test sets.
    3. Build and train the LightGBM model with tuned hyperparameters:
       n_estimators=100, learning_rate=0.05, max_depth=3, num_leaves=8,
       subsample=0.7, colsample_bytree=0.7, min_child_samples=20,
       min_split_gain=0.01, reg_alpha=0.1, reg_lambda=0.1
    4. Evaluate model performance using R², MAE, MSE, RMSE.
    5. Compute feature importance using SHAP values.
    6. Generate residual distribution and parity plots.
    7. Perform 5-Fold cross-validation and summarize metrics.
    8. Save a comprehensive Word report including metrics, SHAP values, 
       residuals, parity plots, and hyperparameters.

Input:
    - Dataset.xlsx: Contains features
      ['Passengers', 'Speed', 'Route_Length_km', 'Number_of_Bus_Stops',
       'Trip_Duration_Minutes', 'Trip_Temprature'] and target 'Trip_Energy_kw'.

Output:
    - Trained LightGBM model
    - Evaluation metrics (train-test and 5-Fold CV)
    - Feature importance (SHAP)
    - Residual and parity plots (PNG)
    - Word report summarizing model performance and details
"""

import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
import shap
import time
import datetime
import lightgbm as lgb
from docx import Document
from docx.shared import Inches
from sklearn.model_selection import train_test_split, KFold
from sklearn.preprocessing import StandardScaler
from sklearn.metrics import mean_squared_error, r2_score, mean_absolute_error
import numpy as np

# Load the dataset
ds = pd.read_excel("Dataset.xlsx")

# Define input features (X) and target variable (y)
feature_names = ['Passengers', 'Speed', 'Route_Length_km', 'Number_of_Bus_Stops',
                 'Trip_Duration_Minutes', 'Trip_Temprature']
X = ds[feature_names]
y = ds['Trip_Energy_kw']

# Split the dataset into training and testing sets (80% training, 20% testing)
X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2, random_state=42)

# Scale the input features
scaler = StandardScaler()
X_train_scaled = scaler.fit_transform(X_train)
X_test_scaled = scaler.transform(X_test)

# LightGBM Model Configuration
lgb_model = lgb.LGBMRegressor(
    n_estimators=100,           # number of trees
    learning_rate=0.05,         # small learning rate for better generalization
    max_depth=3,                # shallow trees to avoid overfitting
    num_leaves=8,               # max leaves for max_depth=3 (2^3)
    subsample=0.7,              # stochastic bagging
    colsample_bytree=0.7,       # feature subsampling
    min_child_samples=20,       # minimum samples per leaf
    min_split_gain=0.01,        # require minimum gain to make split
    reg_alpha=0.1,              # L1 regularization
    reg_lambda=0.1,             # L2 regularization
    random_state=42,
    objective='regression'      # explicitly set regression objective
)


# --- K-Fold Validation ---
kf = KFold(n_splits=5, shuffle=True, random_state=42)
r2_list, mae_list, mse_list, rmse_list = [], [], [], []

for train_index, val_index in kf.split(X_train_scaled):
    X_tr, X_val = X_train_scaled[train_index], X_train_scaled[val_index]
    y_tr, y_val = y_train.iloc[train_index], y_train.iloc[val_index]
    
    lgb_model.fit(X_tr, y_tr)
    y_val_pred = lgb_model.predict(X_val)
    
    r2_list.append(r2_score(y_val, y_val_pred))
    mae_list.append(mean_absolute_error(y_val, y_val_pred))
    mse_list.append(mean_squared_error(y_val, y_val_pred))
    rmse_list.append(mean_squared_error(y_val, y_val_pred, squared=False))

# Train on full training set
start_train = time.time()
lgb_model.fit(X_train_scaled, y_train)
end_train = time.time()
training_time = end_train - start_train

# Predict on test set
start_test = time.time()
y_pred = lgb_model.predict(X_test_scaled)
end_test = time.time()
testing_time = end_test - start_test

# Evaluate final model
mse = mean_squared_error(y_test, y_pred)
r2 = r2_score(y_test, y_pred)
mae = mean_absolute_error(y_test, y_pred)
rmse = mean_squared_error(y_test, y_pred, squared=False)

# SHAP Feature Importance
explainer = shap.Explainer(lgb_model, X_train_scaled)
shap_values = explainer(X_test_scaled)
mean_abs_shap_values = abs(shap_values.values).mean(axis=0)
importance_df = pd.DataFrame({'Feature': feature_names, 'Mean Absolute SHAP Value': mean_abs_shap_values})
importance_df = importance_df.sort_values(by='Mean Absolute SHAP Value', ascending=False)

# Residuals
residuals = y_test - y_pred
mean_resid = residuals.mean()
var_ratio = residuals.var() / y_test.var()
skewness = residuals.skew()

# ----------------- Save Word Report -----------------
doc = Document()
doc.add_heading("LightGBM Model", level=1)
doc.add_paragraph(f"Date: {datetime.datetime.now().strftime('%d-%m-%Y %H:%M')}\n\n")

# 1️⃣ Model Performance Metrics
doc.add_heading("1️⃣ Model Performance Metrics", level=2)
table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Metric'
hdr_cells[1].text = 'Value'

metrics = {
    "R-Squared (R2)": r2,
    "Mean Absolute Error (MAE)": mae,
    "Mean Squared Error (MSE)": mse,
    "Root Mean Squared Error (RMSE)": rmse,
    "Training Time (seconds)": training_time,
    "Testing Time (seconds)": testing_time
}
for metric, value in metrics.items():
    row_cells = table.add_row().cells
    row_cells[0].text = metric
    row_cells[1].text = str(round(value, 4))

# 1a️⃣ K-Fold Validation Metrics
doc.add_heading("K-Fold Validation Metrics (5-fold)", level=2)
table = doc.add_table(rows=1, cols=5)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Model'
hdr_cells[1].text = 'R² (mean ± std)'
hdr_cells[2].text = 'MAE (mean ± std)'
hdr_cells[3].text = 'MSE (mean ± std)'
hdr_cells[4].text = 'RMSE (mean ± std)'

row_cells = table.add_row().cells
row_cells[0].text = 'LightGBM'
row_cells[1].text = f"{np.mean(r2_list):.3f} ± {np.std(r2_list):.3f}"
row_cells[2].text = f"{np.mean(mae_list):.3f} ± {np.std(mae_list):.3f}"
row_cells[3].text = f"{np.mean(mse_list):.3f} ± {np.std(mse_list):.3f}"
row_cells[4].text = f"{np.mean(rmse_list):.3f} ± {np.std(rmse_list):.3f}"

# 2️⃣ Hyperparameters Table
doc.add_heading("2️⃣ Model Hyperparameters", level=2)
table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Model'
hdr_cells[1].text = 'Key Hyperparameters'

row_cells = table.add_row().cells
row_cells[0].text = 'LightGBM'
row_cells[1].text = str(lgb_model.get_params())

# 3️⃣ Feature Importance
doc.add_heading("3️⃣ Feature Importance (Mean Absolute SHAP Values)", level=2)
table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Feature'
hdr_cells[1].text = 'Mean Absolute SHAP Value'

for index, row in importance_df.iterrows():
    row_cells = table.add_row().cells
    row_cells[0].text = row['Feature']
    row_cells[1].text = str(round(row['Mean Absolute SHAP Value'], 4))

# 4️⃣ Residual Distribution Plot
plt.figure(figsize=(10, 6), dpi=300)
sns.kdeplot(residuals, fill=True, color="#1a80bb")
sns.despine(top=True, right=True)
plt.grid(True, which='both', linestyle='--', linewidth=0.5, alpha=0.7)
plt.xlabel("Residuals", fontsize=14, fontname='Times New Roman')
plt.ylabel("Density", fontsize=14, fontname='Times New Roman')
#plt.text(
 #   0.05, 0.95,
  #  f"Mean Residual = {mean_resid:.3f}\nHomoscedasticity = {var_ratio:.3f}\nSkewness = {skewness:.3f}",
   # transform=plt.gca().transAxes, fontname='Times New Roman', fontsize=11,
    #verticalalignment='top', bbox=dict(boxstyle="round,pad=0.3", edgecolor="gray", facecolor="white", alpha=0.8)
#)
residual_plot_filename = "Residual_LightGBM.png"
plt.savefig(residual_plot_filename, dpi=300, bbox_inches='tight')
plt.close()
doc.add_heading("4️⃣ Residual Distribution Plot (KDE)", level=2)
doc.add_picture(residual_plot_filename, width=Inches(5.5))

# Residual Diagnostics Table
doc.add_heading("Residual Diagnostics", level=3)
table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Diagnostic'
hdr_cells[1].text = 'Value'
diagnostics = {"Mean Residual": mean_resid, "Homoscedasticity": var_ratio, "Skewness": skewness}
for diag, value in diagnostics.items():
    row_cells = table.add_row().cells
    row_cells[0].text = diag
    row_cells[1].text = str(round(value, 4))

# 5️⃣ Parity Plot
plt.figure(figsize=(10, 6), dpi=300)
sns.scatterplot(x=y_test, y=y_pred, alpha=0.5, color="#ffb255", edgecolor=None)
plt.plot([y_test.min(), y_test.max()], [y_test.min(), y_test.max()], linestyle="--", color="#000000", lw=2)
sns.despine(top=True, right=True)
plt.grid(True, which="both", linestyle="--", linewidth=0.5, alpha=0.6)
plt.xlabel("Observed Trip Energy (kWh)", fontsize=14, fontname='Times New Roman')
plt.ylabel("Predicted Trip Energy (kWh)", fontsize=14, fontname='Times New Roman')
#plt.text(
 #   0.05, 0.95,
  #  f"R² = {r2:.3f}\nMSE = {mse:.2f}\nRMSE = {rmse:.2f}\nMAE = {mae:.2f}",
   ##verticalalignment='top', bbox=dict(boxstyle="round,pad=0.3", edgecolor="gray", facecolor="white", alpha=0.8)
#)
parity_plot_filename = "Parity_LightGBM.png"
plt.savefig(parity_plot_filename, dpi=300, bbox_inches='tight')
plt.close()
doc.add_heading("5️⃣ Parity Plot (Predicted vs Observed)", level=2)
doc.add_picture(parity_plot_filename, width=Inches(5.5))

# Save Word Document
doc.save("2025_Report_LightGBM_KFold_Full.docx")
print("\n✅ Report saved successfully as '2025_Report_LightGBM_KFold_Full.docx' 🎯")
