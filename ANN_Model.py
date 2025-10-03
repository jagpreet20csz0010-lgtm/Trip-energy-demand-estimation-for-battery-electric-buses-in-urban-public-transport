"""
Title: Trip Energy Demand Estimation – Artificial Neural Network (ANN) Model
Author: Jagpreet Singh
Date: 2025-10-03

Description:
    This script trains and evaluates an Artificial Neural Network (ANN) model 
    to estimate trip energy demand (Trip_Energy_kw) of battery electric buses 
    in urban public transport systems.

Key Steps:
    1. Load and preprocess the dataset (Dataset.xlsx).
    2. Standardize features and split into train/test sets.
    3. Build, compile, and train the ANN model (1 hidden layer with 6 neurons).
    4. Evaluate model performance using R², MAE, MSE, RMSE.
    5. Compute feature importance using SHAP values.
    6. Extract input-to-hidden weights and biases, and construct ANN equation.
    7. Generate residual distribution and parity plots.
    8. Perform 5-Fold cross-validation and summarize metrics.
    9. Save a comprehensive Word report including metrics, feature importance,
       weights, biases, ANN equation, and plots.

Input:
    - 2025_Factorial_Energy_Demand_Dataset.xlsx: Contains features
      ['Passengers', 'Speed', 'Route_Length_km', 'Number_of_Bus_Stops',
       'Trip_Duration_Minutes', 'Trip_Temprature'] and target 'Trip_Energy_kw'.

Output:
    - Trained ANN model
    - Evaluation metrics (train-test and 5-Fold CV)
    - Feature importance (SHAP)
    - Residual and parity plots (PNG)
    - Weights and biases of hidden layer
    - Word report summarizing model performance and details
"""


import pandas as pd
import numpy as np
import tensorflow as tf
from tensorflow import keras
import matplotlib.pyplot as plt
import seaborn as sns
import time
import datetime
import shap
from docx import Document
from docx.shared import Inches
from sklearn.model_selection import train_test_split, KFold
from sklearn.preprocessing import StandardScaler
from sklearn.metrics import mean_absolute_error, mean_squared_error, r2_score

# ---------------------- Helper: create ANN model ----------------------
def create_ann(input_dim):
    tf.random.set_seed(42)
    model = keras.Sequential([
        keras.layers.InputLayer(input_shape=(input_dim,)),
        keras.layers.Dense(6, activation='relu', use_bias=True),
        keras.layers.Dense(1)
    ])
    model.compile(optimizer='adam', loss='mse')
    return model

# ---------------------- Load dataset ----------------------
ds = pd.read_excel("Dataset.xlsx")

# Define input features (X) and target variable (y)
feature_names = ['Passengers', 'Speed', 'Route_Length_km', 'Number_of_Bus_Stops',
                 'Trip_Duration_Minutes', 'Trip_Temprature']
X = ds[feature_names]
y = ds['Trip_Energy_kw']

# ---------------------- Train-test split (original flow preserved) ----------------------
X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2, random_state=42)

# Standardize the input features for the final model
scaler = StandardScaler()
X_train_scaled = scaler.fit_transform(X_train)
X_test_scaled = scaler.transform(X_test)

# ---------------------- Build, compile and train final ANN model (original) ----------------------
ann_model = create_ann(X_train_scaled.shape[1])

start_train = time.time()
ann_model.fit(
    X_train_scaled, y_train,
    validation_split=0.2,   # hold out 20% for validation
    epochs=100,
    batch_size=32,
    verbose=1,
    callbacks=[
        tf.keras.callbacks.EarlyStopping(
            monitor='val_loss',
            patience=10,          # stop if no improvement in 10 epochs
            restore_best_weights=True
        )
    ]
)
end_train = time.time()
training_time = end_train - start_train

# Predict on Test Data
start_test = time.time()
y_pred = ann_model.predict(X_test_scaled).flatten()
end_test = time.time()
testing_time = end_test - start_test

# Evaluate the Model (final)
mae = mean_absolute_error(y_test, y_pred)
mse = mean_squared_error(y_test, y_pred)
rmse = np.sqrt(mse)
r2 = r2_score(y_test, y_pred)

# ---------------------- Feature Importance using SHAP (final model) ----------------------
# Note: For some TF versions SHAP may produce warnings; keeping original approach
explainer = shap.Explainer(ann_model, X_train_scaled)
shap_values = explainer(X_test_scaled)

# Compute mean absolute SHAP values
feature_importance = np.abs(shap_values.values).mean(axis=0)
shap_importance_df = pd.DataFrame({'Feature': feature_names, 'SHAP Importance': feature_importance})
shap_importance_df = shap_importance_df.sort_values(by='SHAP Importance', ascending=False)

# ---------------------- Extract Weights and Biases (input-to-hidden) ----------------------
# Dense(6) is layer index 1 (since InputLayer is layer 0)
input_to_hidden_weights, hidden_layer_biases = ann_model.layers[1].get_weights()
weights_df = pd.DataFrame(input_to_hidden_weights, index=feature_names,
                          columns=[f'H(1:{i+1})' for i in range(input_to_hidden_weights.shape[1])])

# Get output layer (last Dense layer)
output_layer = ann_model.layers[-1]   # safer than using index 2 directly
Output_weights, Output_biases = output_layer.get_weights()

# Create names for hidden neurons based on number of weights coming into output
hidden_neurons = [f'H(1:{i+1})' for i in range(Output_weights.shape[0])]

# Build equation string
equation = "Output = "
equation += f"{round(float(Output_biases[0]), 3)} (Bias for Output Neuron) + "
for i, (w_row, neuron) in enumerate(zip(Output_weights, hidden_neurons)):
    equation += f"{round(float(w_row[0]), 3)} * {neuron}"
    if i < len(Output_weights) - 1:
        equation += " + "


# ---------------------- Residuals & Residual Plot (original) ----------------------
residuals = y_test - y_pred

plt.figure(figsize=(10, 6), dpi=300)
sns.kdeplot(residuals, fill=True, color="#1a80bb")
sns.despine(top=True, right=True)
plt.grid(True, which='both', linestyle='--', linewidth=0.5, alpha=0.7)
plt.xlabel("Residuals", fontsize=14, fontname='Times New Roman')
plt.ylabel("Density", fontsize=14, fontname='Times New Roman')

mean_resid = residuals.mean()
var_ratio = residuals.var() / y_test.var()
skewness = residuals.skew()

# resid_text = (
#     f"Mean Residual = {mean_resid:.3f}\n"
#     f"Homoscedasticity = {var_ratio:.3f}\n"
#     f"Skewness = {skewness:.3f}"
# )
# plt.text(
#     0.05, 0.95, resid_text,
#     transform=plt.gca().transAxes,
#     fontname='Times New Roman',
#     fontsize=11,
#     verticalalignment='top',
#     bbox=dict(boxstyle="round,pad=0.3", edgecolor="gray", facecolor="white", alpha=0.8)
# )

residual_plot_filename = "Residual_ANN.png"
plt.savefig(residual_plot_filename, dpi=300, bbox_inches='tight')
plt.close()

# ---------------------- K-Fold Cross-Validation for ANN (5-fold) ----------------------
kf = KFold(n_splits=5, shuffle=True, random_state=42)
r2_list, mae_list, mse_list, rmse_list = [], [], [], []

X_array = X.values  # original (unscaled) features
y_array = y.values

for fold, (train_idx, val_idx) in enumerate(kf.split(X_array), start=1):
    # Per-fold split and scaling (to avoid leakage)
    X_tr, X_val = X_array[train_idx], X_array[val_idx]
    y_tr, y_val = y_array[train_idx], y_array[val_idx]

    fold_scaler = StandardScaler()
    X_tr_scaled = fold_scaler.fit_transform(X_tr)
    X_val_scaled = fold_scaler.transform(X_val)

    # Recreate model for each fold
    fold_model = create_ann(X_tr_scaled.shape[1])

    # Train (quietly) — keep same epochs & batch size as original
    fold_model.fit(X_tr_scaled, y_tr, epochs=100, batch_size=32, verbose=0)

    # Predict and evaluate
    y_val_pred = fold_model.predict(X_val_scaled).flatten()
    r2_list.append(r2_score(y_val, y_val_pred))
    mae_list.append(mean_absolute_error(y_val, y_val_pred))
    mse_fold = mean_squared_error(y_val, y_val_pred)
    mse_list.append(mse_fold)
    rmse_list.append(np.sqrt(mse_fold))

# Summarize K-Fold metrics (mean ± std)
kf_summary = {
    'R2': (np.mean(r2_list), np.std(r2_list)),
    'MAE': (np.mean(mae_list), np.std(mae_list)),
    'MSE': (np.mean(mse_list), np.std(mse_list)),
    'RMSE': (np.mean(rmse_list), np.std(rmse_list))
}

# ---------------------- Parity Plot (original) ----------------------
plt.figure(figsize=(10, 6), dpi=300)
sns.scatterplot(x=y_test, y=y_pred, alpha=0.5, color="#ffb255", edgecolor=None)
plt.plot([y_test.min(), y_test.max()], [y_test.min(), y_test.max()],
         linestyle="--", color="#000000", lw=2)
sns.despine(top=True, right=True)
plt.grid(True, which='both', linestyle='--', linewidth=0.5, alpha=0.6)
plt.xlabel("Observed Trip Energy (kWh)", fontsize=14, fontname='Times New Roman')
plt.ylabel("Predicted Trip Energy (kWh)", fontsize=14, fontname='Times New Roman')

# metrics_text = (
#     f"R² = {r2:.3f}\n"
#     f"MSE = {mse:.2f}\n"
#     f"RMSE = {rmse:.2f}\n"
#     f"MAE = {mae:.2f}"
# )
# plt.text(
#     0.05, 0.95, metrics_text,
#     transform=plt.gca().transAxes,
#     fontname='Times New Roman',
#     fontsize=11,
#     verticalalignment='top',
#     bbox=dict(boxstyle="round,pad=0.3", edgecolor="gray", facecolor="white", alpha=0.8)
# )

parity_plot_filename = "Parity_ANN.png"
plt.savefig(parity_plot_filename, dpi=300, bbox_inches='tight')
plt.close()

# ---------------------- Save Everything to Word Document ----------------------
doc = Document()
doc.add_heading("Artificial Neural Network Model", level=1)
doc.add_paragraph(f"Date: {datetime.datetime.now().strftime('%d-%m-%Y %H:%M')}\n\n")

# 1️⃣ Model Performance Metrics (final train-test)
doc.add_heading("1️⃣ Model Performance Metrics", level=2)
table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Metric'
hdr_cells[1].text = 'Value'

metrics = {
    "R-Squared (R2)": r2,
    "Mean Absolute Error (MAE)": mae,
    "Mean Squared Error (MSE)": mse,
    "Root Mean Squared Error (RMSE)": rmse,
    "Training Time (seconds)": training_time,
    "Testing Time (seconds)": testing_time
}
for metric, value in metrics.items():
    row_cells = table.add_row().cells
    row_cells[0].text = metric
    row_cells[1].text = str(round(value, 4))

# 1a️⃣ K-Fold CV summary table
doc.add_heading("2️⃣ K-Fold Cross-Validation Performance Metrics (5-Fold)", level=2)
table = doc.add_table(rows=1, cols=5)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = "Model"
hdr[1].text = "R² (mean ± std)"
hdr[2].text = "MAE (mean ± std)"
hdr[3].text = "MSE (mean ± std)"
hdr[4].text = "RMSE (mean ± std)"

row = table.add_row().cells
row[0].text = "ANN"
row[1].text = f"{kf_summary['R2'][0]:.4f} ± {kf_summary['R2'][1]:.4f}"
row[2].text = f"{kf_summary['MAE'][0]:.4f} ± {kf_summary['MAE'][1]:.4f}"
row[3].text = f"{kf_summary['MSE'][0]:.4f} ± {kf_summary['MSE'][1]:.4f}"
row[4].text = f"{kf_summary['RMSE'][0]:.4f} ± {kf_summary['RMSE'][1]:.4f}"

# 2️⃣ Hyperparameters Table (ANN)
doc.add_heading("3️⃣ Model Hyperparameters", level=2)
table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = "Model"
hdr[1].text = "Key Hyper Parameters"

row = table.add_row().cells
row[0].text = "ANN"
ann_params_str = (
    "Architecture: Input -> Dense(6, ReLU) -> Dense(1); "
    "Optimizer=Adam; Loss=MSE; Epochs=100; Batch_size=32; "
    "Activation=ReLU; Random_seed=42"
)
row[1].text = ann_params_str

# 3️⃣ SHAP Feature Importance
doc.add_heading("4️⃣ SHAP Feature Importance", level=2)
table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Feature'
hdr_cells[1].text = 'SHAP Importance'
for _, r in shap_importance_df.iterrows():
    rc = table.add_row().cells
    rc[0].text = r['Feature']
    rc[1].text = str(round(r['SHAP Importance'], 4))

# 4️⃣ Weights and Biases from Input-to-Hidden Layer
doc.add_heading("5️⃣ Weights and Biases from Input-to-Hidden Layer", level=2)
table = doc.add_table(rows=1 + len(feature_names) + 1, cols=weights_df.shape[1] + 1)  # extra col for feature names
table.style = 'Table Grid'
# headers
hdr_cells = table.rows[0].cells
hdr_cells[0].text = "Feature / Bias"
for i, col in enumerate(weights_df.columns):
    hdr_cells[i + 1].text = col
# bias row
row_cells = table.rows[1].cells
row_cells[0].text = "Bias"
for i, val in enumerate(hidden_layer_biases):
    row_cells[i + 1].text = str(round(val, 4))
# feature rows
for row_idx, (feature_name, weight_values) in enumerate(weights_df.iterrows(), start=2):
    rc = table.rows[row_idx].cells
    rc[0].text = feature_name
    for col_idx, val in enumerate(weight_values):
        rc[col_idx + 1].text = str(round(val, 4))

# 5️⃣ Final ANN Equation
doc.add_heading("6️⃣ Final ANN Equation", level=2)
doc.add_paragraph(equation)

# 6️⃣ Residual Distribution Plot
doc.add_heading("7️⃣ Residual Distribution Plot (KDE)", level=2)
doc.add_picture(residual_plot_filename, width=Inches(5.5))

# Residual Diagnostics Table
doc.add_heading("Residual Diagnostics", level=3)
table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Diagnostic'
hdr_cells[1].text = 'Value'
diagnostics = {"Mean Residual": mean_resid, "Homoscedasticity": var_ratio, "Skewness": skewness}
for diag, val in diagnostics.items():
    row_cells = table.add_row().cells
    row_cells[0].text = diag
    row_cells[1].text = str(round(val, 4))

# 7️⃣ Parity Plot
doc.add_heading("8️⃣ Parity Plot (Predicted vs Observed)", level=2)
doc.add_picture(parity_plot_filename, width=Inches(5.5))

# Save the document
doc.save("2025_Report_ANN_KFold_Full.docx")
print("\n✅ Report saved successfully as '2025_Report_ANN_KFold_Full.docx'")
