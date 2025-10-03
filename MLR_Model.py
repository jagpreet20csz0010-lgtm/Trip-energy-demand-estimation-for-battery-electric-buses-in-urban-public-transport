"""
Title: Trip Energy Demand Estimation – Multiple Linear Regression (MLR) Model
Author: Jagpreet Singh
Date: 2025-10-03

Description:
    This script trains and evaluates a Multiple Linear Regression (MLR) model 
    to estimate trip energy demand (Trip_Energy_kw) of battery electric buses 
    in urban public transport systems.

Key Steps:
    1. Load and preprocess the dataset (Dataset.xlsx).
    2. Standardize features and split into train/test sets.
    3. Train the MLR model on training data.
    4. Evaluate model performance using R², MAE, MSE, RMSE.
    5. Compute feature importance using SHAP values.
    6. Generate residual distribution and parity plots.
    7. Perform 5-Fold cross-validation and summarize metrics.
    8. Save a comprehensive Word report including metrics, plots, 
       feature importance, and the final MLR equation.

Input:
    - Dataset.xlsx: Contains features 
      ['Passengers', 'Speed', 'Route_Length_km', 'Number_of_Bus_Stops', 
       'Trip_Duration_Minutes', 'Trip_Temprature'] and target 'Trip_Energy_kw'.

Output:
    - Trained MLR model
    - Evaluation metrics
    - Residual and parity plots (PNG)
    - Word report summarizing model performance and feature importance
"""

import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
import shap
import time
import datetime
from docx import Document
from docx.shared import Inches
from sklearn.model_selection import train_test_split, KFold
from sklearn.preprocessing import StandardScaler
from sklearn.linear_model import LinearRegression
from sklearn.metrics import mean_squared_error, r2_score, mean_absolute_error

# ------------------- Load dataset -------------------
ds = pd.read_excel("Dataset.xlsx")
feature_names = ['Passengers', 'Speed', 'Route_Length_km', 'Number_of_Bus_Stops', 
                 'Trip_Duration_Minutes', 'Trip_Temprature']
X = ds[feature_names]
y = ds['Trip_Energy_kw']

# ------------------- Split and scale dataset -------------------
X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2, random_state=42)
scaler = StandardScaler()
X_train_scaled = scaler.fit_transform(X_train)
X_test_scaled = scaler.transform(X_test)

# ------------------- Train final MLR model -------------------
mlr_model = LinearRegression()
start_train = time.time()
mlr_model.fit(X_train_scaled, y_train)
end_train = time.time()
training_time = end_train - start_train

start_test = time.time()
y_pred = mlr_model.predict(X_test_scaled)
end_test = time.time()
testing_time = end_test - start_test

# ------------------- Evaluate final model -------------------
mse = mean_squared_error(y_test, y_pred)
r2 = r2_score(y_test, y_pred)
mae = mean_absolute_error(y_test, y_pred)
rmse = mean_squared_error(y_test, y_pred, squared=False)

# ------------------- SHAP Feature Importance -------------------
explainer = shap.Explainer(mlr_model, X_train_scaled)
shap_values = explainer(X_test_scaled)
mean_abs_shap_values = abs(shap_values.values).mean(axis=0)
importance_df = pd.DataFrame({'Feature': feature_names, 'Mean Absolute SHAP Value': mean_abs_shap_values})
importance_df = importance_df.sort_values(by='Mean Absolute SHAP Value', ascending=False)

# ------------------- Construct MLR equation -------------------
mlr_equation = f"Trip_Energy_kw = {mlr_model.intercept_:.4f}"
for feature, coef in zip(feature_names, mlr_model.coef_):
    mlr_equation += f" + ({coef:.4f} * {feature})"

# ------------------- Residuals -------------------
residuals = y_test - y_pred

# ------------------- Residual Distribution Plot -------------------
plt.figure(figsize=(10, 6), dpi=300)
sns.kdeplot(residuals, fill=True, color="#1a80bb")
sns.despine(top=True, right=True)
plt.grid(True, which='both', linestyle='--', linewidth=0.5, alpha=0.7)
plt.xlabel("Residuals", fontsize=14, fontname='Times New Roman')
plt.ylabel("Density", fontsize=14, fontname='Times New Roman')

mean_resid = residuals.mean()
var_ratio = residuals.var() / y_test.var()
skewness = residuals.skew()



residual_plot_filename = "Residual_MLR.png"
plt.savefig(residual_plot_filename, dpi=300, bbox_inches='tight')
plt.close()

# ------------------- Parity Plot -------------------
plt.figure(figsize=(10, 6), dpi=300)
sns.scatterplot(x=y_test, y=y_pred, alpha=0.5, color="#ffb255", edgecolor=None)
plt.plot([y_test.min(), y_test.max()], [y_test.min(), y_test.max()], linestyle="--", color="#000000", lw=2)
sns.despine(top=True, right=True)
plt.grid(True, which="both", linestyle="--", linewidth=0.5, alpha=0.6)
plt.xlabel("Observed Trip Energy (kWh)", fontsize=14, fontname='Times New Roman')
plt.ylabel("Predicted Trip Energy (kWh)", fontsize=14, fontname='Times New Roman')



parity_plot_filename = "Parity_MLR.png"
plt.savefig(parity_plot_filename, dpi=300, bbox_inches='tight')
plt.close()

# ------------------- K-Fold Validation -------------------
kf = KFold(n_splits=5, shuffle=True, random_state=42)
r2_list, mse_list, rmse_list, mae_list = [], [], [], []

X_scaled = scaler.fit_transform(X)  # scale full dataset
y_array = y.values

for train_idx, test_idx in kf.split(X_scaled):
    X_train_kf, X_test_kf = X_scaled[train_idx], X_scaled[test_idx]
    y_train_kf, y_test_kf = y_array[train_idx], y_array[test_idx]
    model_kf = LinearRegression()
    model_kf.fit(X_train_kf, y_train_kf)
    y_pred_kf = model_kf.predict(X_test_kf)
    r2_list.append(r2_score(y_test_kf, y_pred_kf))
    mse_list.append(mean_squared_error(y_test_kf, y_pred_kf))
    rmse_list.append(mean_squared_error(y_test_kf, y_pred_kf, squared=False))
    mae_list.append(mean_absolute_error(y_test_kf, y_pred_kf))

# Compute mean ± std
metrics_summary = {
    'R2': (np.mean(r2_list), np.std(r2_list)),
    'MSE': (np.mean(mse_list), np.std(mse_list)),
    'RMSE': (np.mean(rmse_list), np.std(rmse_list)),
    'MAE': (np.mean(mae_list), np.std(mae_list))
}

# ------------------- Save to Word Document -------------------
doc = Document()
doc.add_heading("Multiple Linear Regression Model", level=1)
doc.add_paragraph(f"Date: {datetime.datetime.now().strftime('%d-%m-%Y %H:%M')}\n\n")

# 1️⃣ Model Performance Metrics (Original Train-Test)
doc.add_heading("1️⃣ Model Performance Metrics (Train-Test Split)", level=2)
table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Metric'
hdr_cells[1].text = 'Value'

metrics = {
    "Intercept": mlr_model.intercept_,
    "R-Squared (R2)": r2,
    "Mean Absolute Error (MAE)": mae,
    "Mean Squared Error (MSE)": mse,
    "Root Mean Squared Error (RMSE)": rmse,
    "Training Time (seconds)": training_time,
    "Testing Time (seconds)": testing_time
}

for metric, value in metrics.items():
    row_cells = table.add_row().cells
    row_cells[0].text = metric
    row_cells[1].text = str(round(value, 4))

# 2️⃣ K-Fold CV Metrics
doc.add_heading("2️⃣ K-Fold Cross-Validation Performance Metrics", level=2)
table = doc.add_table(rows=1, cols=5)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Model'
hdr_cells[1].text = 'R² (mean ± std)'
hdr_cells[2].text = 'MAE (mean ± std)'
hdr_cells[3].text = 'MSE (mean ± std)'
hdr_cells[4].text = 'RMSE (mean ± std)'

row_cells = table.add_row().cells
row_cells[0].text = "MLR"
row_cells[1].text = f"{metrics_summary['R2'][0]:.3f} ± {metrics_summary['R2'][1]:.3f}"
row_cells[2].text = f"{metrics_summary['MAE'][0]:.3f} ± {metrics_summary['MAE'][1]:.3f}"
row_cells[3].text = f"{metrics_summary['MSE'][0]:.3f} ± {metrics_summary['MSE'][1]:.3f}"
row_cells[4].text = f"{metrics_summary['RMSE'][0]:.3f} ± {metrics_summary['RMSE'][1]:.3f}"

# 3️⃣ Hyperparameters Table
doc.add_heading("3️⃣ Model Hyperparameters", level=2)
table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Model'
hdr_cells[1].text = 'Key Hyper Parameters'

row_cells = table.add_row().cells
row_cells[0].text = "MLR"
row_cells[1].text = "Ordinary Least Squares, fit_intercept=True, features standardized, evaluated using 5-Fold CV"

# 4️⃣ Feature Importance
doc.add_heading("4️⃣ Feature Importance (Mean Absolute SHAP Values)", level=2)
table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Feature'
hdr_cells[1].text = 'Mean Absolute SHAP Value'

for index, row in importance_df.iterrows():
    row_cells = table.add_row().cells
    row_cells[0].text = row['Feature']
    row_cells[1].text = str(round(row['Mean Absolute SHAP Value'], 4))

# 5️⃣ Final MLR Equation
doc.add_heading("5️⃣ Final Multiple Linear Regression Model Equation", level=2)
doc.add_paragraph(mlr_equation)

# 6️⃣ Residual Distribution Plot
doc.add_heading("6️⃣ Residual Distribution Plot (KDE)", level=2)
doc.add_picture(residual_plot_filename, width=Inches(5.5))

# Residual Diagnostics Table
doc.add_heading("Residual Diagnostics", level=3)
table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Diagnostic'
hdr_cells[1].text = 'Value'

diagnostics = {
    "Mean Residual": mean_resid,
    "Homoscedasticity": var_ratio,
    "Skewness": skewness
}

for diag, value in diagnostics.items():
    row_cells = table.add_row().cells
    row_cells[0].text = diag
    row_cells[1].text = str(round(value, 4))

# 7️⃣ Parity Plot
doc.add_heading("7️⃣ Parity Plot (Predicted vs Observed)", level=2)
doc.add_picture(parity_plot_filename, width=Inches(5.5))

# ------------------- Save Document -------------------
doc.save("2025_Report_MLR_KFold_Full.docx")
print("\n✅ Report saved successfully as '2025_Report_MLR_KFold_Full.docx'")
